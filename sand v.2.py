import math, io
from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, Circle

# ==========================
# Функция расчёта геометрии и массы
# ==========================
def compute(L, W, phi_deg, rho):
    m = 1.0 / math.tan(math.radians(phi_deg))
    Hmax = min(L/(2*m), W/(2*m))
    Ltop = max(L - 2*m*Hmax, 0)
    Wtop = max(W - 2*m*Hmax, 0)
    Atop = Ltop * Wtop
    Abot = L * W
    if Atop > 0:
        V = (Hmax/3.0) * (Abot + Atop + math.sqrt(Abot*Atop))
    else:
        V = (Hmax/3.0) * Abot
    M = rho * V
    return {"m": m, "Hmax": Hmax, "Ltop": Ltop, "Wtop": Wtop, "Atop": Atop, "Abot": Abot, "V": V, "M": M}

# ==========================
# Функция генерации отчёта с автоматическим масштабированием графиков
# ==========================
def generate_sand_sites_report_scaled(sites_params, doc_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)
    doc.add_heading('Вместимость площадок для временного хранения песка — схемы и расчёты', level=1)
    doc.add_paragraph('Общая схема и расчёты для всех площадок.')

    # --- Общая схема откоса ---
    fig, ax = plt.subplots(figsize=(6,4))
    H_example = 1.0
    m_example = 1.6
    dx = m_example * H_example
    ax.plot([-dx, 0, dx], [0, H_example, 0], marker='o', linewidth=2)
    ax.plot([-dx-0.5, dx+0.5], [0,0], linewidth=1.0)
    ax.annotate("", xy=(0,H_example), xytext=(0.4,H_example-0.1), arrowprops=dict(arrowstyle="->"))
    ax.text(0.5, H_example-0.15, "H (высота)", va="center", fontsize=9)
    ax.annotate("", xy=(dx,0), xytext=(dx+0.5,0.15), arrowprops=dict(arrowstyle="->"))
    ax.text(dx+0.6, 0.15, r"$\varphi$ (угол откоса)", fontsize=9, va="bottom")
    ax.annotate("", xy=(-dx/2, H_example/2), xytext=(-2.5,0.6), arrowprops=dict(arrowstyle="->"))
    ax.text(-2.5, 0.6, r"$m=\cot\varphi=\Delta/H$", fontsize=9)
    plt.subplots_adjust(bottom=0.28)
    ax.text(0.0, -0.35, "Схема (вид в профиль): H — высота насыпи; Δ = m·H — горизонтальный отход откоса; φ — угол откоса; m = 1/tan(φ)",
            transform=ax.transAxes, fontsize=9, ha='left')
    ax.set_xlim(-4,4); ax.set_ylim(-1.0,1.6); ax.set_aspect('equal', adjustable='box'); ax.axis('off')

    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    doc.add_heading('Общая схема (профиль откоса)', level=2)
    doc.add_paragraph('Схема показывает профиль откоса, определение высоты H и коэффициента m = cot(φ).')
    doc.add_picture(buf, width=Inches(6))
    buf.close()

    # --- По каждой площадке ---
    for idx, params in enumerate(sites_params, start=1):
        L = params["L"]
        W = params["W"]
        phi = params["phi"]
        rho = params["rho"]
        name = params["name"]

        res = compute(L,W,phi,rho)
        m = res["m"]
        H = res["Hmax"]
        Ltop = res["Ltop"]
        Wtop = res["Wtop"]
        Atop = res["Atop"]
        Abot = res["Abot"]
        V = res["V"]
        M = res["M"]

        # --- масштабирование графиков ---
        max_plot_size = 10
        scale = max(L, W, H*2) / max_plot_size

        # --- План ---
        fig, ax = plt.subplots(figsize=(max(L/scale,6), max(W/scale,4)))
        pad = max(L,W)*0.05
        ax.set_xlim(-pad, L+pad)
        ax.set_ylim(-pad, W+pad)
        base = Rectangle((0,0), L, W, fill=False, linewidth=1.5)
        ax.add_patch(base)
        ax.text(0.5, W+pad, f"План. Base: {int(L)}×{int(W)} м", fontsize=9)
        if Atop > 0.0 and (Ltop>0 and Wtop>0):
            left = (L - Ltop)/2.0
            bottom = (W - Wtop)/2.0
            top_rect = Rectangle((left, bottom), Ltop, Wtop, fill=False, linestyle='--')
            ax.add_patch(top_rect)
            ax.text(left + Ltop/2.0 - 1, bottom + Wtop/2.0, f"Atop={Atop:.1f} м²", fontsize=9)
        else:
            cx = L/2.0; cy = W/2.0
            ax.add_patch(Circle((cx, cy), 0.3))
            ax.text(cx+0.5, cy, "Apex (вершина)", fontsize=9)
        ax.set_title(f"План: {name}", fontsize=9)
        ax.set_aspect('equal', adjustable='box')
        ax.axis('off')

        plan_buf = io.BytesIO()
        fig.savefig(plan_buf, format='png', bbox_inches="tight", dpi=150)
        plt.close(fig)
        plan_buf.seek(0)

        # --- Профиль ---
        fig, ax = plt.subplots(figsize=(max(L/scale,6), max(H*2/scale,3)))
        off = m * H
        if Ltop > 0 and Ltop > 0.001:
            x_pts = [0, off, L-off, L]; y_pts = [0, H, H, 0]
        else:
            cx = L/2.0; x_pts = [0, cx, L]; y_pts = [0, H, 0]
        ax.plot(x_pts, y_pts, marker='o', linewidth=1.5)
        ax.plot([0,L], [0,0], linewidth=1.0)
        xpos = L/2.0
        ax.annotate("H", xy=(xpos, H), xytext=(xpos+1, H/2), arrowprops=dict(arrowstyle="->"))
        ax.annotate(rf"$\Delta = mH = {off:.2f}\ \mathrm{{м}}$", xy=(L/4, H/4), xytext=(L/4, H/4+0.5),
                    arrowprops=dict(arrowstyle="->"))
        ax.set_xlim(-1, L+1); ax.set_ylim(-0.5, H+2)
        ax.set_xlabel("м по базовой длине", fontsize=9)
        ax.set_ylabel("высота, м", fontsize=9)
        ax.grid(True, linewidth=0.3)
        ax.set_title(f"Профиль: {name} — H={H:.3f} м, m={m:.3f}, Atop={Atop:.1f} м²", fontsize=9)

        profile_buf = io.BytesIO()
        fig.savefig(profile_buf, format='png', bbox_inches="tight", dpi=150)
        plt.close(fig)
        profile_buf.seek(0)

        # --- Добавляем в Word ---
        doc.add_page_break()
        doc.add_heading(f'Площадка {idx}: {name}', level=2)
        doc.add_paragraph('План:')
        doc.add_picture(plan_buf, width=Inches(6))
        plan_buf.close()
        doc.add_paragraph('Профиль:')
        doc.add_picture(profile_buf, width=Inches(6))
        profile_buf.close()
        doc.add_paragraph('Расчёты:')
        calc_text = (
            f"L = {L} м, W = {W} м, φ = {phi}°, ρ = {rho} т/м³\n"
            f"m = cot(φ) = {m:.6f}\n"
            f"H_max = {H:.6f} м\n"
            f"L_top = {Ltop:.6f} м, W_top = {Wtop:.6f} м -> A_top = {Atop:.3f} м²\n"
            f"A_bot = {Abot:.3f} м²\n"
            f"Объём V = {V:.3f} м³\n"
            f"Масса M = {M:.3f} т\n"
        )
        for line in calc_text.splitlines():
            doc.add_paragraph(line, style='List Bullet')

    doc.save(doc_path)
    print("Готово — Word сохранён в:", doc_path)
    return doc_path

# ==========================
# Пример использования
# ==========================
sites_params = [
    {"L":30.0, "W":30.0, "phi":32.0, "rho":1.60, "name":"Площадка 30×30, Сухой песок", "H_max_limit":6.0},
    {"L":30.0, "W":30.0, "phi":28.0, "rho":1.85, "name":"Площадка 30×30, Влажный песок"},
    {"L":50.0, "W":50.0, "phi":32.0, "rho":1.60, "name":"Площадка 50×50, Сухой песок"},
    {"L":50.0, "W":50.0, "phi":28.0, "rho":1.85, "name":"Площадка 50×50, Влажный песок"},
]

doc_path = "Площадки_расчёт_и_схемы_масштабированные.docx"
generate_sand_sites_report_scaled(sites_params, doc_path)
